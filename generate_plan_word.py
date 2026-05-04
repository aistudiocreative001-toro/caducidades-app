from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título principal
title = doc.add_heading('CADFZ — APP DE CADUCIDADES', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtítulo
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Plan de Implementación Definitivo\nStack: Next.js 14 + Vercel Blob + Upstash Redis + Tailwind CSS')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 1. ARQUITECTURA
doc.add_heading('1. Arquitectura y Stack Tecnológico', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Capa'
hdr_cells[1].text = 'Tecnología'
hdr_cells[2].text = 'Por qué'

stack = [
    ('Framework', 'Next.js 14 (App Router)', 'SSR/CSR híbrido, Server Actions nativos, despliegue en Vercel.'),
    ('UI / CSS', 'Tailwind CSS + shadcn/ui', 'Material Design limpio, utility-first, rápido.'),
    ('Tablas', 'TanStack Table v8', 'Ordenación, filtrado, selección múltiple.'),
    ('Datos', 'TanStack Query (React Query)', 'Caché inteligente, optimistic updates.'),
    ('Animaciones', 'Framer Motion', 'Entradas escalonadas, modales, tarjetas.'),
    ('Almacenamiento', 'Vercel Blob', 'CSV como base de datos única, lectura/escritura milisegundos.'),
    ('Concurrencia', 'Upstash Redis', 'Lock para escrituras simultáneas + cache JSON parseado (TTL 5 min).'),
    ('CSV', 'PapaParse', 'Robustez con encodings mixtos y formato europeo.'),
    ('Fechas', 'date-fns', 'Manipulación segura y ligera.'),
    ('Iconos', 'lucide-react', 'Conjunto moderno y consistente.')
]

for layer, tech, why in stack:
    row_cells = table.add_row().cells
    row_cells[0].text = layer
    row_cells[1].text = tech
    row_cells[2].text = why

doc.add_paragraph()

# 2. MODELO DE DATOS
doc.add_heading('2. Modelo de Datos (TypeScript)', level=1)
doc.add_paragraph(
    'Cada producto tiene un campo id (UUID v4) invisible para el usuario, pero imprescindible '
    'para editar o borrar sin confusión (el CODIGO se repite para el mismo producto con otra fecha).'
)

fields = [
    'id: string — UUID generado en migración',
    'ubi: string — LR | 3C | CL | AL',
    'codigo: string — Código de barras (NO único global)',
    'sku: string',
    'producto: string — Nombre descriptivo',
    'marca: string',
    'tipo: string — Categoría (Proteína, Comida Fitness, etc.)',
    'coste: number — Coste unitario (parseado de "35,44" → 35.44)',
    'uds: number — Unidades disponibles',
    'costeTotal: number — uds * coste (auto)',
    'fecha: string — ISO 8601: "2026-03-31"',
    'fechaMes: string — "mar-2026" (auto desde fecha)',
    'dias: number — Días hasta caducidad vs hoy (auto)',
    'estado: string — Estado actual (ver tabla abajo)',
    'observaciones: string',
    'tags: string — String de tags separados por coma'
]

for f in fields:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f).font.name = 'Consolas'

doc.add_paragraph()

# 3. ESTADOS
doc.add_heading('3. Estados del Sistema', level=1)
states = [
    ('VIGENTE', 'Calculado', 'dias > 30, fecha >= hoy'),
    ('PRÓXIMO / EN RIESGO', 'Calculado', '0 <= dias <= 30'),
    ('CADUCADO', 'Calculado', 'dias < 0 y no es estado fijo'),
    ('ROTO', 'Fijo', 'Producto dañado físicamente. Nunca recalcula.'),
    ('VENDIDO', 'Fijo', 'Agotado por venta. Nunca recalcula.'),
    ('VENDIDO CADUCADO', 'Fijo', 'Vendido ya caducado (solo desde admin). Nunca recalcula.'),
    ('REGALO CADUCADO', 'Fijo', 'Regalado ya caducado (solo desde admin). Nunca recalcula.'),
    ('MOVIDO', 'Fijo', 'Todo movido desde almacén. Nunca recalcula.'),
    ('MOSTRADOR', 'Mixto', 'Estado manual, pero si fecha < hoy → CADUCADO al recalcular.'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Estado'
hdr2[1].text = 'Tipo'
hdr2[2].text = 'Descripción'

for s, t, d in states:
    r = table2.add_row().cells
    r[0].text = s
    r[1].text = t
    r[2].text = d

doc.add_paragraph()
preg = doc.add_paragraph()
preg.add_run('Regla de recálculo de estado: ').bold = True
preg.add_run('Al cargar datos, si fecha < hoy y el estado NO está en la lista de fijos → se fuerza a CADUCADO.')

doc.add_paragraph()

# 4. ALMACENAMIENTO CSV
doc.add_heading('4. Almacenamiento: CSV en Vercel Blob', level=1)

doc.add_heading('Migración inicial', level=2)
steps = [
    'Script único que lee el CSV local "CADUCIDADES hasta 09_26 incluido - Caducidades.csv".',
    'Salta las 3 primeras líneas (encabezados de Excel).',
    'Detecta encabezados reales: UBI, CODIGO, SKU, PRODUCTO, MARCA, TIPO, COSTE, UDS, COSTE TOTAL, FECHA, FECHA MES, DIAS, ESTADO, OBSERVACIONES, TAGS.',
    'Para cada fila no vacía: genera id (crypto.randomUUID()), parsea coste/costeTotal reemplazando coma por punto, parsea fecha de DD/MM/YYYY a ISO YYYY-MM-DD, recalcula dias, fechaMes y estado inicial.',
    'Escribe CSV limpio a Vercel Blob (caducidades.csv).'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# 5. LECTURA/ESCRITURA
doc.add_heading('5. Lógica de Lectura y Escritura', level=1)

doc.add_paragraph('getProducts() — Server Action de lectura:', style='Heading 3')
read_steps = [
    'Intenta leer cache Redis (clave products:parsed). Si existe y TTL no expiró → devuelve inmediatamente.',
    'Si no en cache: lee Blob, parsea con PapaParse, recalcula dias/estado, guarda en Redis (TTL 300s).',
    'Devuelve JSON al frontend.'
]
for s in read_steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('Server Actions de escritura — CRUD, Vender, Mover:', style='Heading 3')
write_steps = [
    'Intenta adquirir lock Redis (SET lock:caducidades NX EX 10).',
    'Si lock conseguido: lee CSV de Blob → modifica en memoria (RAM) → convierte a CSV → sobrescribe Blob → invalida cache Redis (DEL products:parsed).',
    'Si no consigue lock: espera 100ms, reintenta hasta 5 veces.',
    'Libera lock (expira automático a los 10s como seguridad).'
]
for s in write_steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# 6. CADUCADOS HOY
doc.add_heading('6. Lógica de "Caducados Hoy"', level=1)

doc.add_heading('Al arrancar la app', level=2)
start_steps = [
    'Recorre todos los productos.',
    'Si fecha < hoy Y estado NO es fijo Y estado NO es ya CADUCADO:',
    '   - Cambia estado a CADUCADO.',
    '   - Guarda su id en Redis bajo clave caducados:YYYY-MM-DD (ej: caducados:2026-05-02).',
    'Escribe CSV actualizado con los estados corregidos.'
]
for s in start_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('Modal "Han caducado N productos"', level=2)
modal_steps = [
    'Al entrar al frontend, TanStack Query llama a getProducts().',
    'Si el servidor detectó N productos que pasaron a CADUCADO hoy Y localStorage NO tiene la marca de hoy (caducadosVistos:YYYY-MM-DD) → muestra modal OBLIGATORIO (no se cierra haciendo click fuera).',
    'El modal lista los productos que caducaron, con botón "Aceptar".',
    'Al pulsar Aceptar → guarda marca en localStorage, modal cierra, página refresca datos.',
    'Al día siguiente: nueva fecha, nuevo modal si hay nuevos caducados.'
]
for s in modal_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('Botón "Caducados a fecha" (Panel Admin)', level=2)
doc.add_paragraph(
    'En el panel de administración, un botón que consulta la clave Redis del día actual '
    'y muestra un modal con el listado de productos que hoy pasaron a CADUCADO. '
    'Este listado es temporal (un día).'
)

doc.add_page_break()

# 7. PANTALLAS
doc.add_heading('7. Flujo de Pantallas y Componentes', level=1)

# Dashboard
doc.add_heading('Pantalla 1: Dashboard Principal (/)', level=2)
doc.add_paragraph(
    'Layout: AppBar blanco con título "CadFZ" y botón "Panel de Administración" a la derecha. '
    'Grid de 4 tarjetas (responsive: 1 móvil, 2 tablet, 4 desktop).'
)
doc.add_heading('Especificaciones de tarjeta tienda:', level=3)
tarjeta = [
    'Fondo blanco #FFFFFF, esquinas rounded-2xl (16px).',
    'Borde: 1px #E2E8F0.',
    'Barra lateral izquierda: 4px de color tienda (LR morado #7B1FA2, 3C verde #43A047, CL amarillo #FBC02D, AL naranja #FB8C00).',
    'Sombra: shadow-sm en reposo, shadow-md + translate-y-[-2px] en hover.',
    'Padding: p-6 (1.5rem / 24px).',
    'Contenido vertical: icono circular 48px con fondo 10% opaco del color tienda → nombre tienda (text-xl font-bold) → uds y coste (text-sm slate-500) → badges estado (inline-flex, fondo pastel).',
    'Badges: 🔴 Caducados (#FEE2E2 / #DC2626), 🟠 Próximos (#FFF7ED / #EA580C), 🟢 Vigentes (#ECFDF5 / #059669).'
]
for t in tarjeta:
    doc.add_paragraph(t, style='List Bullet')

# Vista tienda
doc.add_heading('Pantalla 2: Vista por Tienda (/tienda/[ubi])', level=2)

doc.add_heading('A. Tiendas LR, 3C, CL', level=3)
ltienda = [
    'Grid de tarjetas una por categoría (tipo).',
    'Solo productos donde fecha >= hoy Y uds > 0 Y estado NO es fijo final (VENDIDO, ROTO, VENDIDO CADUCADO, REGALO CADUCADO, MOVIDO).',
    'Muestra el producto con menor dias (más urgente) de esa categoría.',
    'Desplegable "Siguiente más próximo": lista vertical sin tarjeta, animación height 0→auto.'
]
for t in ltienda:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('Contenido de tarjeta:')
tarjcont = [
    'Cabecera: nombre categoría (ej: Proteínas).',
    'Metadata: UBI · MARCA · TIPO (ej: LR · Biotech USA · Proteína).',
    'Nombre producto: font-semibold text-lg.',
    'Detalles: 3 uds · 31/03/2026 · 15 días · 106,32 €.',
    'Acción: input cantidad editable (default 1, ancho 60px) + botón Vender.',
    'Validación: cantidad <= uds. Si no, borde rojo.',
    'Optimistic UI: resta inmediatamente. Si nuevasUds == 0: animación scale→0.9, opacity→0, remove del DOM. Estado histórico: VENDIDO.'
]
for t in tarjcont:
    doc.add_paragraph(t, style='List Bullet 2')

doc.add_heading('B. Almacén (AL) — Caso especial', level=3)
almacen = [
    'Mismas tarjetas por categoría, mismas reglas de filtrado.',
    'Colores de prioridad por días del producto mostrado (barras laterales):',
    '   🔴 URGENTE: 0–30 días',
    '   🟠 PRIORITARIO: 30–60 días',
    '   🟡 PROGRAMAR: 60–90 días',
    '   >90 días: sin color especial (gris/verde suave).',
    'Sin botón Vender. Botón "Mover".',
    'Todos los productos de AL se muestran (no hay filtro de 90 días).'
]
for a in almacen:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('Modal "Mover" (Almacén)', level=3)
modalmover = [
    'Título: Mover producto. Producto origen (nombre, uds disponibles, coste).',
    '3 inputs stepper (botones - / + y campo numérico): LR, 3C, CL.',
    'Validación en tiempo real: suma de destinos <= uds origen. Si excede → inputs en rojo, botón Confirmar deshabilitado.',
    'Al confirmar: Lock Redis → lee CSV → modifica AL (uds -= totalMovido) → crea NUEVOS registros en destino (mismos datos excepto ubi y uds asignadas, nuevo id) → si AL queda en uds == 0, estado = MOVIDO → escribe CSV → invalida cache.',
    'AL queda con uds: 0, estado: MOVIDO. Tarjeta desaparece de vista AL.',
    'Los registros nuevos aparecerán en /tienda/LR, /tienda/3C, /tienda/CL con estado recalculado desde su fecha.'
]
for m in modalmover:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# Panel Admin
doc.add_heading('Pantalla 3: Panel de Administración (/admin)', level=2)

doc.add_heading('KPIs Dinámicos (tarjetas por estado)', level=3)
kpis = [
    'Fila de tarjetas informativas (no clickeables), una por cada estado real existente en productos filtrados.',
    'Cada tarjeta muestra: nombre estado, total unidades (suma de uds), coste total (suma de costeTotal).',
    'Colores de fondo suave y texto por estado (ver tabla de colores).',
    'Dinámicos: al usar filtros (búsqueda, tienda, categoría, rango días), TODAS las tarjetas recalculan sus totales solo con productos que cumplen el filtro.',
    'Ejemplo: Filtro "Proteína" → KPIs cambian. Filtro búsqueda "iso" → KPIs cambian. Filtro rango URGENTE (0-30) → KPIs solo con dias <= 30.'
]
for k in kpis:
    doc.add_paragraph(k, style='List Bullet')

doc.add_heading('Filtros de admin', level=3)
filtros = [
    'Búsqueda: input texto libre. Filtra por producto, sku, codigo, marca. Debounce 300ms, mínimo 2 caracteres.',
    'Tienda: select (LR, 3C, CL, AL, Todas).',
    'Categoría: select (lista dinámica desde tipos existentes + Todas).',
    'Rango días: select (TODOS / URGENTE 0-30 / PRÓXIMO 30-60 / PROGRAMAR 60-90).',
    'Los 4 filtros son COMBINABLES y afectan SIMULTÁNEAMENTE a tarjetas KPI + tabla de productos.'
]
for f in filtros:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Tabla de productos', level=3)
tabla = [
    'TanStack Table con columnas: Checkbox, UBI, CÓDIGO, SKU, PRODUCTO, MARCA, TIPO, COSTE, UDS, COSTE TOTAL, FECHA, DIAS, ESTADO, OBSERVACIONES, TAGS, ACCIONES.',
    'Ordenación por defecto: dias ASC.',
    'Estado de fila con badge coloreado.',
    'Fila hover: bg-slate-50.',
    'Acciones por fila: ✏️ Editar (abre drawer) / 🗑️ Eliminar (modal confirmación).',
    'Selección múltiple con checkbox → toolbar inferior para "Eliminar seleccionados" o "Cambiar estado en masa".'
]
for t in tabla:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Botones especiales admin', level=3)
botones = [
    'Nuevo Producto: abre Drawer lateral vacío.',
    'Importar CSV: input file. Al seleccionar CSV, parsea y añade/actualiza registros (match por codigo + fecha para actualizar, crear si no existe).',
    'Exportar a Excel: genera CSV con separador ; y descarga con extensión .csv.',
    'Caducados a fecha: consulta clave Redis del día y muestra modal con listado de productos que hoy pasaron a CADUCADO.'
]
for b in botones:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Drawer: Crear/Editar Producto', level=3)
drawer = [
    'Slide-in desde derecha. Ancho max-w-md desktop, full-screen móvil. Backdrop: bg-black/20.',
    'Campos editables: UBI, CÓDIGO, SKU, PRODUCTO, MARCA, TIPO (dropdown), COSTE, UDS, FECHA (datepicker), ESTADO (searchable select, permite crear libre), OBSERVACIONES, TAGS.',
    'Campos auto (solo lectura): DIAS, FECHA MES, COSTE TOTAL (uds * coste).',
    'Botones: Eliminar (izquierda) y Guardar (derecha, azul primario #1565C0).',
    'Al guardar: bloquea input, muestra loading spinner, optimistic update en tabla.'
]
for d in drawer:
    doc.add_paragraph(d, style='List Bullet')

doc.add_page_break()

# 8. COLORES DEFINITIVOS
doc.add_heading('8. Colores Definitivos', level=1)

doc.add_heading('Tiendas (badge, barra lateral, icono)', level=2)
tcolores = [
    ('LR', '#7B1FA2', 'Morado'),
    ('3C', '#43A047', 'Verde'),
    ('CL', '#FBC02D', 'Amarillo (texto oscuro #1E293B)'),
    ('AL', '#FB8C00', 'Naranja')
]
table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
h3 = table3.rows[0].cells
h3[0].text = 'Tienda'
h3[1].text = 'Hex'
h3[2].text = 'Nombre'
for t, h, n in tcolores:
    r = table3.add_row().cells
    r[0].text = t
    r[1].text = h
    r[2].text = n

doc.add_heading('Estados (badge admin KPI)', level=2)
escolores = [
    ('VIGENTE', '#ECFDF5', '#059669'),
    ('PRÓXIMO / EN RIESGO', '#FFF7ED', '#EA580C'),
    ('CADUCADO', '#FEF2F2', '#DC2626'),
    ('ROTO', '#FEF3C7', '#D97706'),
    ('VENDIDO', '#F1F5F9', '#475569'),
    ('VENDIDO CADUCADO', '#FEE2E2', '#B91C1C'),
    ('REGALO CADUCADO', '#FCE7F3', '#BE185D'),
    ('MOVIDO', '#E0E7FF', '#4F46E5'),
    ('MOSTRADOR', '#E0F2FE', '#0284C7')
]
table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'
h4 = table4.rows[0].cells
h4[0].text = 'Estado'
h4[1].text = 'Fondo'
h4[2].text = 'Texto'
for e, f, t in escolores:
    r = table4.add_row().cells
    r[0].text = e
    r[1].text = f
    r[2].text = t

doc.add_heading('Urgencia (barras laterales AL)', level=2)
urgcol = [
    ('0-30 días', '#E53935', 'Rojo'),
    ('30-60 días', '#FB8C00', 'Naranja'),
    ('60-90 días', '#FBC02D', 'Amarillo'),
    ('>90 días', '#43A047', 'Verde')
]
table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Light Grid Accent 1'
h5 = table5.rows[0].cells
h5[0].text = 'Rango'
h5[1].text = 'Hex'
h5[2].text = 'Color'
for rg, hx, cl in urgcol:
    r = table5.add_row().cells
    r[0].text = rg
    r[1].text = hx
    r[2].text = cl

doc.add_heading('Tokens globales', level=2)
tokens = [
    ('Body background', '#F3F4F6'),
    ('Tarjeta / Superficie', '#FFFFFF'),
    ('Texto primario', '#0F172A'),
    ('Texto secundario/label', '#64748B'),
    ('Borde/divisor', '#E2E8F0'),
    ('Botón primario', '#1565C0'),
    ('Botón destructivo', '#DC2626'),
    ('Éxito', '#059669'),
    ('Advertencia', '#D97706')
]
table6 = doc.add_table(rows=1, cols=2)
table6.style = 'Light Grid Accent 1'
h6 = table6.rows[0].cells
h6[0].text = 'Token'
h6[1].text = 'Valor'
for tk, vl in tokens:
    r = table6.add_row().cells
    r[0].text = tk
    r[1].text = vl

doc.add_page_break()

# 9. ANIMACIONES
doc.add_heading('9. Animaciones (Framer Motion)', level=1)
animaciones = [
    ('Tarjetas dashboard entrada', 'staggerChildren 0.05, opacity 0→1, y: 20→0, duration 0.3'),
    ('Tarjeta hover', 'y: -2, boxShadow: 0 4px 6px -1px rgba(0,0,0,0.1)'),
    ('Tarjeta vendida/movida (uds→0)', 'scale: 1→0.9, opacity: 1→0, height: auto→0, remove DOM'),
    ('Modal abrir', 'Backdrop opacity 0→1, Modal scale 0.95→1, y: 20→0'),
    ('Drawer abrir', 'x: 100%→0, duration 0.3, ease: easeOut'),
    ('Desplegable siguiente', 'AnimatePresence, height 0→auto, opacity 0→1'),
    ('Notificación toast', 'x: 100→0, opacity 0→1, auto-dismiss 3s'),
    ('Botón click', 'scale: 0.98, duration 0.1'),
    ('Fila tabla hover', 'backgroundColor transition 150ms'),
    ('Input focus', 'border-color 150ms, box-shadow 0 0 0 3px rgba(37,99,235,0.1)')
]
table7 = doc.add_table(rows=1, cols=2)
table7.style = 'Light Grid Accent 1'
h7 = table7.rows[0].cells
h7[0].text = 'Interacción'
h7[1].text = 'Implementación'
for ani, impl in animaciones:
    r = table7.add_row().cells
    r[0].text = ani
    r[1].text = impl

doc.add_page_break()

# 10. SERVER ACTIONS
doc.add_heading('10. Server Actions (Next.js 14 App Router)', level=1)
sa = [
    ('getProducts()', '-', 'Product[]', 'No (lectura Redis/Blob)'),
    ('getProductsByUbi(ubi)', 'ubi: string', 'Record<string, Product[]>', 'No'),
    ('venderProducto(id, cantidad)', 'id: string, cantidad: number', '{ ok, product }', 'Sí (lock Redis)'),
    ('moverProducto(id, destinos)', 'id: string, destinos: { LR?, 3C?, CL? }', '{ ok }', 'Sí'),
    ('createProduct(data)', 'Omit<Product, "id"|"dias"|"fechaMes"|"costeTotal">', 'Product', 'Sí'),
    ('updateProduct(id, data)', 'id: string, Partial<Product>', 'Product', 'Sí'),
    ('deleteProducts(ids)', 'ids: string[]', '{ ok, deleted }', 'Sí'),
    ('importCSV(file)', 'file: File (multipart)', '{ created, updated }', 'Sí'),
    ('exportCSV()', '-', 'Blob (descarga)', 'No'),
    ('getCaducadosDeHoy()', '-', 'Product[]', 'No')
]
table8 = doc.add_table(rows=1, cols=4)
table8.style = 'Light Grid Accent 1'
h8 = table8.rows[0].cells
h8[0].text = 'Action'
h8[1].text = 'Parámetros'
h8[2].text = 'Retorno'
h8[3].text = 'Lock'
for a, p, r, l in sa:
    row = table8.add_row().cells
    row[0].text = a
    row[1].text = p
    row[2].text = r
    row[3].text = l

doc.add_paragraph()
err = doc.add_paragraph()
err.add_run('Manejo de errores: ').bold = True
err.add_run('Si lock no disponible tras 5 reintentos → retorna 429. Si cantidad > uds → 422. Catch general → 500 con log servidor.')

doc.add_page_break()

# 11. VARIABLES ENTORNO
doc.add_heading('11. Variables de Entorno (.env.local)', level=1)
envs = [
    'BLOB_READ_WRITE_TOKEN=         # Vercel Blob (read/write)',
    'UPSTASH_REDIS_REST_URL=        # Ej: https://xxx.upstash.io',
    'UPSTASH_REDIS_REST_TOKEN=      # Token Upstash REST'
]
for e in envs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(e).font.name = 'Consolas'

doc.add_paragraph('Configuración en Vercel: Dashboard → Proyecto → Settings → Environment Variables.')

# 12. DEPENDENCIAS
doc.add_heading('12. Dependencias (package.json)', level=1)

doc.add_paragraph('dependencies:').bold = True
deps = [
    'next ^14.2.0',
    'react ^18.3.0',
    'react-dom ^18.3.0',
    '@tanstack/react-table ^8.17.0',
    '@tanstack/react-query ^5.40.0',
    '@vercel/blob ^0.23.0',
    '@upstash/redis ^1.32.0',
    'papaparse ^5.4.1',
    'date-fns ^3.6.0',
    'lucide-react ^0.379.0',
    'framer-motion ^11.2.0',
    'clsx ^2.1.0',
    'tailwind-merge ^2.3.0'
]
for d in deps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(d).font.name = 'Consolas'

doc.add_paragraph('devDependencies:').bold = True
devs = [
    'typescript ^5.4.0',
    '@types/node ^20.0.0',
    '@types/react ^18.3.0',
    '@types/papaparse ^5.3.14',
    'tailwindcss ^3.4.0',
    'autoprefixer ^10.4.0',
    'postcss ^8.4.0'
]
for d in devs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(d).font.name = 'Consolas'

doc.add_page_break()

# 13. ROADMAP
doc.add_heading('13. Roadmap de Construcción', level=1)
roadmap = [
    ('Fase 1', 'Setup', 'Crear proyecto Next.js, Tailwind, instalar deps, configurar types, variables de entorno.'),
    ('Fase 2', 'Migración', 'Script migrate.ts. Parsear CSV local, añadir UUIDs, escribir a Blob.'),
    ('Fase 3', 'Backend', 'Server Actions: getProducts, getProductsByUbi, vender, mover, CRUD. Lock Redis + cache.'),
    ('Fase 4', 'Dashboard', 'Página / con 4 tarjetas tienda, navegación, animaciones.'),
    ('Fase 5', 'Vista Tienda', 'Página /tienda/[ubi], tarjetas categoría, botón Vender + Mover, optimistic UI.'),
    ('Fase 6', 'Admin', 'Página /admin, KPIs dinámicos, tabla TanStack, filtros, drawer crear/editar.'),
    ('Fase 7', 'Import/Export', 'Botones importar CSV, exportar Excel.'),
    ('Fase 8', 'Caducados hoy', 'Lógica recálculo, modal inicial, botón Caducados a fecha.'),
    ('Fase 9', 'Polish', 'Responsive móvil, animaciones Framer Motion, toast notificaciones, loading states.'),
    ('Fase 10', 'Deploy', 'Push a GitHub, importar en Vercel, configurar env vars.')
]
table9 = doc.add_table(rows=1, cols=3)
table9.style = 'Light Grid Accent 1'
h9 = table9.rows[0].cells
h9[0].text = 'Fase'
h9[1].text = 'Nombre'
h9[2].text = 'Descripción'
for f, n, d in roadmap:
    r = table9.add_row().cells
    r[0].text = f
    r[1].text = n
    r[2].text = d

# 14. DECISIONES
doc.add_heading('14. Decisiones Críticas Confirmadas', level=1)
decs = [
    'Colores tienda: LR morado #7B1FA2, 3C verde #43A047, CL amarillo #FBC02D, AL naranja #FB8C00.',
    '"Mover" desde AL crea registros NUEVOS en destino (no suma a existentes).',
    'AL usa colores de prioridad por días (0-30 rojo, 30-60 naranja, 60-90 amarillo).',
    'Si uds AL llega a 0 tras mover, estado = MOVIDO.',
    'En vista tienda (LR/3C/CL) solo se muestran productos vigentes (fecha >= hoy).',
    'Panel admin muestra TODOS los registros, incluyendo vendidos/caducados/rotos.',
    'KPIs admin informativos, no clickeables. Se filtran con selects arriba.',
    'KPIs se recalculan dinámicamente con cualquier filtro (búsqueda, tienda, cat, rango días).',
    'Estado auto-recalculado al arrancar: vigentes con fecha pasada → CADUCADO.',
    'Modal "Han caducado N productos" aparece una vez por sesión (localStorage).',
    'Botón "Caducados a fecha" en admin consulta clave Redis del día actual.',
    'Lock Redis para escrituras simultáneas, máximo 5 reintentos.'
]
for d in decs:
    doc.add_paragraph(d, style='List Number')

# Guardar
doc.save(r'C:\Users\ivhea\Desktop\CadFZ OpenCode\Plan_CadFZ_Caducidades.docx')
print('Documento Word generado exitosamente: Plan_CadFZ_Caducidades.docx')
